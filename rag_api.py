import os
import glob
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain.schema import Document as LC_Document
from docx import Document as DocxDocument  # for .docx support
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

def load_docx_as_documents(file_path):
    doc = DocxDocument(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    return [LC_Document(page_content=full_text)]

def load_and_index_documents(directory="documents", index_path="faiss_index"):
    file_paths = glob.glob(f"{directory}/*.txt")
    pdf_paths = glob.glob(f"{directory}/*.pdf")
    docx_paths = glob.glob(f"{directory}/*.docx")
    documents = []

    for path in file_paths:
        loader = TextLoader(path)
        documents.extend(loader.load())

    for path in pdf_paths:
        loader = PyPDFLoader(path)
        documents.extend(loader.load())

     # Load DOCX files
    for path in docx_paths:
        documents.extend(load_docx_as_documents(path))

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    split_docs = splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    db = FAISS.from_documents(split_docs, embeddings)
    db.save_local(index_path)

def setup_rag(index_path="faiss_index"):
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

    client = OpenAI(
        base_url="https://models.github.ai/inference",
        api_key=GITHUB_TOKEN,
    )

    class GitHubRAG:
        def __init__(self, retriever, client):
            self.retriever = retriever
            self.client = client

        def run(self, question):
            docs = self.retriever.invoke(question)
            context = "\n".join([doc.page_content for doc in docs]) or "No documents found."

            if not context.strip():
                return "I couldn't find any relevant information in the documents to answer your question."

            prompt = f"""
You are CHATDOC, a RAG assistant designed to answer questions strictly using the provided document context.

Your job is to:
- Explain, summarize, or answer **any question about the content, sections, structure, or metadata** of the document (e.g., author, supervisor, dedication, purpose, table of contents, chapter details, etc.).
- If the user asks about a specific section, search the context and extract or summarize the most relevant part.
- If the user asks who the author, supervisor, or document is dedicated to, extract that information directly from the context if it exists.
- If the exact answer is not found, try to rephrase the user’s question to something similar and check again. In such a case, start your response with:  
  *Do you mean ‘[rephrased question]’?*
- If the information is still not found, respond with:  
  *I could not find relevant information in the documents.*

Use only the context below to answer the question. Do not make up answers.

Context:
{context}

Question:
{question}
"""

            response = self.client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that strictly answers based on provided context."},
                    {"role": "user", "content": prompt}
                ],
                model="openai/gpt-4o",
                temperature=0.7,
                max_tokens=1024
            )
            return response.choices[0].message.content

    return GitHubRAG(retriever, client)
